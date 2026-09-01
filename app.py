import os
import io
import re
import zipfile
from datetime import datetime
from itertools import zip_longest
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

TEMPLATE_PATH = "letter_template.docx"

def get_formatted_today():
    day = datetime.now().day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return datetime.now().strftime(f"{day}{suffix} %B %Y")

def format_date_str(date_obj):
    day = date_obj.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return date_obj.strftime(f"{day}{suffix} %B %Y")

def professional_title_case(text):
    """Formats research titles professionally by keeping minor prepositions and conjunctions lowercase."""
    lowercase_words = {
        'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 
        'to', 'from', 'by', 'with', 'in', 'of', 'over', 'under', 'via'
    }
    words = text.strip().split()
    if not words:
        return ""
    
    formatted = []
    for i, word in enumerate(words):
        word_clean = word.lower()
        # Always capitalize first and last word; keep minor words lowercase inside
        if i == 0 or i == len(words) - 1 or word_clean not in lowercase_words:
            formatted.append(word.capitalize())
        else:
            formatted.append(word_clean)
            
    return " ".join(formatted)

def extract_org_name(address_text):
    """Extracts the top line of the address to serve as the Granting/Host Organization Name."""
    lines = [line.strip() for line in address_text.split("\n") if line.strip()]
    if lines:
        first_line = lines[0]
        return first_line.rstrip(",.").strip()
    return ""

def generate_single_docx(template_path, replacements, address_text, location_text, student_name_raw, research_title_raw):
    """Generates a single in-memory Word document stream for a given address and location."""
    doc = Document(template_path)
    address_lines = [line.strip() for line in address_text.split("\n") if line.strip()]

    org_name = extract_org_name(address_text)

    current_replacements = replacements.copy()
    current_replacements["{LOCATIONS}"] = location_text
    current_replacements["{ORGANIZATION_NAME}"] = org_name

    def safe_replace(paragraph):
        # Identify Paragraph 4 specifically via its intro phrase
        is_paragraph_4 = "the purpose of this letter is" in paragraph.text.lower()

        # Fix missing 'to introduce' in paragraph 4
        if "the purpose of this letter is {STUDENT_NAME}" in paragraph.text:
            for run in paragraph.runs:
                if "is {STUDENT_NAME}" in run.text:
                    run.text = run.text.replace("is {STUDENT_NAME}", "is to introduce {STUDENT_NAME}")
                elif "is " in run.text and "{STUDENT_NAME}" not in run.text:
                    run.text = run.text.replace("is ", "is to introduce ")

        # Address Block handling: Bold + Arial 12pt + 1.5 Line Spacing
        if "{ADDRESS}" in paragraph.text:
            if address_lines:
                paragraph.text = ""
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.5
                
                for i, line in enumerate(address_lines):
                    run = paragraph.add_run(line)
                    run.font.name = "Arial"
                    run.font.size = Pt(12)
                    run.bold = True
                    if i < len(address_lines) - 1:
                        paragraph.add_run("\n")
            return

        # Perform placeholder replacements
        for key, value in current_replacements.items():
            if key in paragraph.text:
                # Decide case formatting dynamically
                if is_paragraph_4:
                    if key == "{STUDENT_NAME}":
                        target_value = student_name_raw.strip().title()
                    elif key == "{RESEARCH_TITLE}":
                        target_value = professional_title_case(research_title_raw)
                    else:
                        target_value = value
                else:
                    # Every other location (including REF / Subject line) stays UPPERCASE
                    if key == "{STUDENT_NAME}":
                        target_value = student_name_raw.strip().upper()
                    elif key == "{RESEARCH_TITLE}":
                        target_value = research_title_raw.strip().upper()
                    else:
                        target_value = value

                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, target_value)

        # Remove extra double spacing inside (REG. NO.  XXX) patterns
        if "(REG. NO.  " in paragraph.text:
            for run in paragraph.runs:
                if "(REG. NO.  " in run.text:
                    run.text = run.text.replace("(REG. NO.  ", "(REG. NO. ")

    # Process paragraphs & tables
    for paragraph in list(doc.paragraphs):
        safe_replace(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    safe_replace(paragraph)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_letter():
    if not os.path.exists(TEMPLATE_PATH):
        return render_template('index.html', error=f"Template file '{TEMPLATE_PATH}' not found.")

    # 1. Parse Dates
    try:
        start_dt = datetime.strptime(request.form.get("start_date"), "%Y-%m-%d")
        end_dt = datetime.strptime(request.form.get("end_date"), "%Y-%m-%d")
    except (ValueError, TypeError):
        return render_template('index.html', error="Invalid date format provided.")

    duration_days = (end_dt - start_dt).days

    if duration_days <= 0:
        return render_template('index.html', error="Research end date must be after the start date.")

    # 2. Check Degree Program (PhD vs Master)
    program_text = request.form.get("program", "").strip().lower()
    short_form_text = request.form.get("short_form", "").strip().lower()

    is_phd = "phd" in program_text or "ph.d" in program_text or "phd" in short_form_text or "ph.d" in short_form_text or "doctor" in program_text

    if is_phd:
        min_days = 30
        max_days = 180  # 1 to 6 months
        degree_label = "PhD"
        month_limit = "1 and 6 months (30 to 180 days)"
    else:
        min_days = 30
        max_days = 90   # 1 to 3 months
        degree_label = "Master's"
        month_limit = "1 and 3 months (30 to 90 days)"

    if duration_days < min_days or duration_days > max_days:
        return render_template(
            'index.html', 
            error=f"Out of Bound! Research duration for {degree_label} students must be between {month_limit}. Your entry was {duration_days} days."
        )

    formatted_start = format_date_str(start_dt)
    formatted_end = format_date_str(end_dt)
    today_date = get_formatted_today()

    raw_reg = request.form.get("reg_no", "").strip()
    address_raw = request.form.get("address", "").strip()
    locations_raw = request.form.get("locations", "").strip()

    student_name_raw = request.form.get("student_name", "").strip()
    research_title_raw = request.form.get("research_title", "").strip()

    replacements = {
        "{DATE}": today_date,
        "{STUDENT_NAME}": student_name_raw.upper(),
        "{REG_NO}": raw_reg,
        "{REG_NO.}": raw_reg,
        "{PROGRAM}": request.form.get("program", "").strip(),
        "{SHORT_FORM}": request.form.get("short_form", "").strip(),
        "{RESEARCH_TITLE}": research_title_raw.upper(),
        "{START_DATE}": formatted_start,
        "{END_DATE}": formatted_end,
    }

    # Support flexible delimiters: '--', '---', or newline dashes
    addresses = [a.strip() for a in re.split(r'-{2,}', address_raw) if a.strip()]
    locations = [l.strip() for l in re.split(r'-{2,}', locations_raw) if l.strip()]

    clean_reg = raw_reg.replace("/", "_").replace("\\", "_").replace(" ", "_")

    # Single Document Case
    if len(addresses) == 1:
        loc_text = locations[0] if locations else ""
        docx_stream = generate_single_docx(TEMPLATE_PATH, replacements, addresses[0], loc_text, student_name_raw, research_title_raw)
        output_filename = f"Clearance_Letter_{clean_reg}.docx"
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Batch Case
    zip_stream = io.BytesIO()
    doc_counter = 1
    default_location = locations[0] if locations else ""

    with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for addr, loc in zip_longest(addresses, locations, fillvalue=default_location):
            if not addr:
                continue
            
            docx_stream = generate_single_docx(TEMPLATE_PATH, replacements, addr, loc, student_name_raw, research_title_raw)
            filename_in_zip = f"Clearance_Letter_{clean_reg}_Target_{doc_counter}.docx"
            zip_file.writestr(filename_in_zip, docx_stream.getvalue())
            doc_counter += 1

    zip_stream.seek(0)
    zip_filename = f"Clearance_Letters_{clean_reg}_Batch.zip"

    return send_file(
        zip_stream,
        as_attachment=True,
        download_name=zip_filename,
        mimetype="application/zip"
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)